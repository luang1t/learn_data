from docx import Document
from docx.shared import Pt

#Criando novo documento
doc = Document()

#Título Principal
doc.add_heading('Guia de estudos: Lógica de Análise de Dados com Pandas', 0)

#Subtítulo com detalhes do projeto
p = doc.add_paragraph()
p.add_run('Projeto: ').bold = True
p.add_run('Análise Exploratória da Copa do Mundo 2022\n')
p.add_run('Data: ').bold = True
p.add_run('29/12/2025\n')

#---- parte 1 ----
doc.add_heading('1. O Papel da Analista de Dados', level = 1)
p = doc.add_paragraph()
p.add_run('Antes de escrever o código, o analista precisa entender o "comportamento" dos dados.').italic = True
p.add_run('No nosso projeto, não começamos criando gráficos complexos; começamos investigando a estrutura básica da tabela.\n')
p.add_run('Conceito Chave: ').bold = True
p.add_run('Data Literacy (Alfabetização de Dados) é a capacidade de ler, trabalhar, analisar e argumentar com dados.\n')

#--- parte 2 ----
doc.add_heading('2. A Ferramenta de Raio-X: .describe()', level = 1)
doc.add_paragraph('O método .describe() é a nossa primeira linha de defesa para entender o dataset desconhecido. Ele nos dá um resumo estético das colunas numéricas.')

doc.add_heading('O Caso da Posse de Bola (total_possession)', level=2)
doc.add_paragraph('Ao rodarmos o describe nesse coluna, observamos: ')

#Lista com bullets
items_secao2= [
    "Média (Mean) ≈ 100: Isso foi um grande indício. Em estatística de futebol, a posse é percentual. A soma resulta em 100%.",
    "Desvio Padrão (std) ≈ 0.24: Indica que os dados quase não variam em relação à média (são consistentes)."
]

for item in items_secao2:
    doc.add_paragraph(item, style='List Bullet')

doc.add_paragraph('Conclusão: A soma da posse é sempre 100%, com pequenas variações apenas por arredondamento.')

#--- parte 3 ---
doc.add_heading('3. Estrutura de Dados: "Wide" vs. "Long"', level = 1)
doc.add_paragraph('Desafio: O dataset é "por partida", mas queríamos analisar "por time".')

items_secao3 = [
    "O problema O Brasil não aparece sempre na mesma coluna (ora é home, ora é away).",
    "Solução Lógica: Olhar o Brasil como Mandante (max) + Olhar o Brasil como Visitanten(max) -> Comparar os dois."
]

for item in items_secao3:
    doc.add_paragraph(item,style='List Bullet')

#--- parte 4 ---
doc.add_heading('4. Dicionário de Termos Técnicos', level=1)
termos = [
    ("DataFrame:", "A tabela inteira."),
    ("Series:", "Uma única coluna (ex: df['coluna'])."),
    ("Std (Desvio Padrão):", "Medida de dispersão (o quanto foge da média)."),
    ("Query:", "Ato de filtrar linhas com uma regra.")
]

for termo, definicao in termos:
    p = doc.add_paragraph()
    p.add_run(termo).bold = True
    p.add_run(definicao)

#--- parte 5 ---
doc.add_heading('5. Perguntas para o Notebook', level=1)
doc.add_paragraph('Use estas perguntas para testar seu conhecimento: ')
perguntas = [
    "Explique a relação entre 'std' baixo e previsibilidade dos dados.",
    "Por que não posso filtrar apenas a coluna 'home_team' para ver todos os jogos do Brasil?",
    "O que o método describe() revelou que evitou um erro de interpretação?"
]

for pergunta in perguntas:
    doc.add_paragraph(pergunta, style='List Number')

#Salva o aquivo

file_name = 'Guia_Analise_Copa.docx'
doc.save(file_name)
print(f"Arquivo '{file_name}' criado com SUCESSO! Pode baixar e subir no NotebookLM")   